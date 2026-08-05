"""Export reports to PDF, DOCX, Markdown, TXT, and JSON."""
from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Any

from models.schemas import ResearchState
from utils.helpers import slugify


def to_markdown(state: ResearchState) -> str:
    parts = [f"# {state.query}", ""]
    if state.summary:
        parts += [f"## Executive Summary\n\n{state.summary}", ""]
    parts += [state.improved_report or state.report, ""]
    if state.citations:
        parts += ["## References", ""] + state.citations + [""]
    if state.interview:
        parts.append("## Interview Questions\n")
        for i, item in enumerate(state.interview, 1):
            parts.append(f"**Q{i}.** {item.get('question','')}\n\n*Answer:* {item.get('answer','')}\n")
    if state.quiz:
        parts.append("\n## Quiz\n")
        for i, q in enumerate(state.quiz, 1):
            parts.append(f"**Q{i}.** {q.get('question','')}")
            for j, opt in enumerate(q.get("options", [])):
                parts.append(f"  - {'ABCD'[j]}. {opt}")
            parts.append(f"  *Answer:* {'ABCD'[q.get('answer',0)]}\n")
    return "\n".join(parts)


def to_text(state: ResearchState) -> str:
    md = to_markdown(state)
    # Strip basic markdown markers
    for ch in ("#", "*", "_", "`"):
        md = md.replace(ch, "")
    return md


def to_json(state: ResearchState) -> str:
    return json.dumps(state.model_dump(), indent=2, default=str)


def to_pdf(state: ResearchState) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted,
    )
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14, alignment=TA_LEFT)
    pre = ParagraphStyle("Pre", parent=styles["Code"], fontSize=8, leading=11)

    story: list[Any] = [Paragraph(state.query, h1), Spacer(1, 12)]
    if state.summary:
        story += [Paragraph("Executive Summary", h2), Paragraph(state.summary, body), Spacer(1, 10)]

    # Render report markdown headings as Paragraphs
    report = state.improved_report or state.report
    for line in report.split("\n"):
        if line.startswith("## "):
            story += [Spacer(1, 6), Paragraph(line[3:].strip(), h2)]
        elif line.startswith("# "):
            story += [Spacer(1, 6), Paragraph(line[2:].strip(), h1)]
        elif line.strip():
            story.append(Paragraph(line, body))
        else:
            story.append(Spacer(1, 4))

    if state.citations:
        story += [PageBreak(), Paragraph("References", h1)]
        for c in state.citations:
            story.append(Paragraph(c.replace("\n", "<br/>"), body))
            story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


def to_docx(state: ResearchState) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(state.query, level=0)
    if state.summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(state.summary)

    report = state.improved_report or state.report
    for line in report.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip():
            doc.add_paragraph(line)

    if state.citations:
        doc.add_heading("References", level=1)
        for c in state.citations:
            doc.add_paragraph(c)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export(state: ResearchState, fmt: str) -> tuple[str, bytes]:
    """Return (filename, bytes) for the requested format."""
    base = slugify(state.query, max_len=40)
    fmt = fmt.lower()
    if fmt == "pdf":
        return f"{base}.pdf", to_pdf(state)
    if fmt == "docx":
        return f"{base}.docx", to_docx(state)
    if fmt == "json":
        return f"{base}.json", to_json(state).encode("utf-8")
    if fmt == "txt":
        return f"{base}.txt", to_text(state).encode("utf-8")
    # default markdown
    return f"{base}.md", to_markdown(state).encode("utf-8")
